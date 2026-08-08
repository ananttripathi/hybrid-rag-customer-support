"""Generates an editable Comparative_Analysis_Report.docx from the same real evaluation artefacts
used by scripts/make_comparative_report_pdf.py."""
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "Comparative_Analysis_Report.docx"

v1_summary = pd.read_csv("v1_summary_metrics.csv")
v2_impact = pd.read_csv("v1_vs_v2_finetuning_impact.csv")
comp_summary = pd.read_csv("Comparative_Results_Summary.csv")
with open("training_reproducibility.json") as f:
    train_info = json.load(f)
with open("outputs.json") as f:
    outputs = json.load(f)
with open("prep_config.json") as f:
    prep_info = json.load(f)

router_df = pd.read_csv("router_evaluation.csv")
full_fmt = router_df["router_parsed_ok"].mean() * 100
full_exact = router_df["exact_match"].mean() * 100
full_fuzzy = router_df["fuzzy_match"].mean() * 100
adv = router_df[router_df["is_adversarial"]]
adv_fmt = adv["router_parsed_ok"].mean() * 100 if len(adv) else float("nan")
adv_exact = adv["exact_match"].mean() * 100 if len(adv) else float("nan")
adv_fuzzy = adv["fuzzy_match"].mean() * 100 if len(adv) else float("nan")

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x0b, 0x3d, 0x91)

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x16, 0x32, 0x4f)

def p(text_runs):
    para = doc.add_paragraph()
    if isinstance(text_runs, str):
        text_runs = [(text_runs, False, False)]
    for text, bold, mono in text_runs:
        r = para.add_run(text)
        r.bold = bold
        if mono:
            r.font.name = "Courier New"
    return para

def bullets(items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        if isinstance(item, str):
            item = [(item, False, False)]
        for text, bold, mono in item:
            r = para.add_run(text)
            r.bold = bold
            if mono:
                r.font.name = "Courier New"

def mono_block(text):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()

title("Comparative Analysis Report")
subtitle("Hybrid RAG & Fine-Tuning for Customer Support")

# ---------------- 1. Compare all versions ----------------
h1("1. Compare All Solution Versions (Task 4.5.1)")
p(f"All three architectures — Baseline (zero-shot {prep_info['MODEL_ID']}), Solution V1 (Naive RAG: "
  f"raw-query retrieval + generation), and Solution V2 (Hybrid RAG: LoRA-fine-tuned intent router "
  f"driving retrieval + the same generation model) — were evaluated under an identical configuration: "
  f"same MODEL_ID, same embedding model (sentence-transformers/all-MiniLM-L6-v2), same ChromaDB index, "
  f"same held-out test partition, and identical deterministic decoding "
  f"(do_sample=False, temperature=None).")

h2("ROUGE / BLEU — SOP-grounded, all 3 architectures")
table(["System", "ROUGE-1", "ROUGE-L", "BLEU"],
      [[r["System"], f"{r['ROUGE-1']:.3f}", f"{r['ROUGE-L']:.3f}", f"{r['BLEU']:.3f}"]
       for _, r in comp_summary.iterrows()])

h2("Format Adherence & Intent Accuracy — Solution V2 Router")
table(["Split", "Format Adherence", "Exact Match", "Fuzzy Match"], [
    ["Full held-out test split (zero leakage)", f"{full_fmt:.1f}%", f"{full_exact:.1f}%", f"{full_fuzzy:.1f}%"],
    ["Adversarial subset (regex-filtered)", f"{adv_fmt:.1f}%", f"{adv_exact:.1f}%", f"{adv_fuzzy:.1f}%"],
])
p("The adversarial subset isolates queries containing hedging/sentiment language (still, never, "
  "terrible, frustrated, ...) — exactly the noisy phrasing that defeats naive semantic retrieval and "
  "motivates the fine-tuned router in the first place. In this sample it came back empty (0/149): the "
  "Bitext dataset is largely template-generated and rarely uses this kind of sentiment language, so "
  "the adversarial-accuracy columns above are reported as N/A rather than a misleadingly perfect/"
  "undefined score.")

h2("Baseline vs Solution V1 — Retrieval's Independent Impact (Task 3.4.1)")
table(list(v1_summary.columns), v1_summary.round(3).values.tolist())
p("Sign convention: for every metric except Hallucination Frequency, a positive Change(%) is an "
  "improvement; for Hallucination Frequency, a negative Change(%) is the improvement (fewer "
  "hallucinations).")

h2("Solution V1 vs Solution V2 — Fine-Tuning's Independent Impact (Task 4.4.2)")
table(list(v2_impact.columns), v2_impact.round(3).values.tolist())
p("Because the generation model is held identical between V1 and V2 (Section 2 of this report), this "
  "delta isolates fine-tuning's contribution from retrieval's — the two effects measured "
  "independently in Tasks 3.4.1 and 4.4.2 together decompose the full system's end-to-end improvement.")

doc.add_page_break()

# ---------------- 2. Reproducibility ----------------
h1("2. Configuration & Reproducibility")
p("A single consistent configuration was used across every notebook, per the assignment's "
  "Implementation Guidance:")
table(["Setting", "Value"], [
    ["MODEL_ID", train_info["MODEL_ID"]],
    ["Embedding model", "sentence-transformers/all-MiniLM-L6-v2"],
    ["LoRA target modules", ", ".join(train_info["lora_config"]["target_modules"])],
    ["LoRA r / alpha / dropout", f"{train_info['lora_config']['r']} / {train_info['lora_config']['lora_alpha']} / {train_info['lora_config']['lora_dropout']}"],
    ["Learning rate / batch size", f"{train_info['learning_rate']} / {train_info['batch_size']}"],
    ["Training steps (max / actual)", f"{train_info['max_steps']} / {train_info['actual_steps']}"],
    ["Early stopping", f"patience={train_info['patience']}, triggered={train_info['early_stopped']}"],
    ["Best validation loss (step)", f"{train_info['best_val_loss']:.4f} (step {train_info['best_step']})"],
    ["Train / Valid / Test rows", f"{prep_info['n_train']} / {prep_info['n_valid']} / {prep_info['n_test']}"],
    ["Execution device", f"{train_info['device']} — {train_info['precision']}"],
    ["Random seed", str(train_info["seed"])],
])

try:
    doc.add_picture("training_curves.png", width=Inches(6.0))
    cap = doc.add_paragraph()
    r = cap.add_run("Figure: training vs validation loss for the LoRA intent router (Notebook 6).")
    r.italic = True
except Exception:
    pass

doc.add_page_break()

# ---------------- 3. End-to-end example ----------------
h1("3. End-to-End Example (Same Query, All 3 Systems)")
p([("Test query: ", True, False), (f"“{outputs['test_query']}”", False, False)])
p([("Ground truth SOP rule: ", True, False), (outputs["ground_truth"], False, False)])
h2("Baseline")
mono_block(outputs.get("baseline_output", "")[:600])
h2("Naive RAG (V1)")
mono_block(outputs.get("naive_rag_output", "")[:600])
p([("Retrieved document: ", False, False), (outputs.get("naive_rag_retrieved_doc", ""), False, True)])
h2("Hybrid RAG (V2)")
mono_block(f"Router output: {outputs.get('hybrid_rag_router_output', '')}")
mono_block(outputs.get("hybrid_rag_output", "")[:600])

doc.add_page_break()

# ---------------- 4. Findings ----------------
h1("4. Findings, Limitations, and Recommendations (Task 4.5.2)")

h2("Key Findings")
bullets([
    [("Retrieval alone is necessary but not sufficient. ", True, False),
     ("Naive RAG raised Format Adherence and reduced Hallucination Frequency relative to the Baseline, "
      "but top-1 similarity search against the raw, noisy customer query is fragile — the "
      "RAG_Implementation notebook (Task 3.2.3) shows this directly: sarcastic/ambiguous phrasing can "
      "still retrieve the wrong or a diluted SOP.", False, False)],
    [("Fine-tuning closes the retrieval-quality gap. ", True, False),
     ("Routing on the LoRA-extracted structured intent (Task 4.3.1) instead of the raw query "
      "consistently selects the intended SOP document, which is what Task 4.4.2's V1-vs-V2 comparison "
      "isolates and quantifies.", False, False)],
    [("Hallucination frequency ", True, False),
     ("is the most business-relevant metric in this domain: fabricated refund windows or invented "
      "escalation paths carry direct customer-trust and compliance risk. Both retrieval and "
      "fine-tuning move this metric in the right direction independently.", False, False)],
    [("Consistency Rate held at 100% ", True, False),
     ("throughout, confirming deterministic (do_sample=False, temperature=None) inference was "
      "correctly configured everywhere — a prerequisite for trustworthy A/B comparison between systems.", False, False)],
])

h2("Limitations")
bullets([
    "Executed on Apple Silicon (MPS) rather than a CUDA T4 GPU: 4-bit quantisation "
    "(bitsandbytes/QLoRA) was unavailable, so LoRA was run at higher precision (fp16 inference / "
    "bf16 + gradient checkpointing for training) instead — documented in the Project Proposal (1.4.1, 1.4.3).",
    "Dataset was downsampled to 1,500 rows (within the assignment's 1,000-5,000 range) and evaluation "
    "loops were run on bounded subsamples of the held-out test split, to keep multi-hour LLM-generation "
    "loops tractable on local hardware rather than a T4 GPU — the assignment explicitly permits this "
    "('df_test.sample(50) is acceptable', 'df_test.head(50) if time-constrained').",
    "Hallucination detection uses a heuristic (unsupported numeric claims + ungrounded terminology "
    "relative to the SOP reference) rather than human annotation or an LLM-judge — adequate for "
    "directional comparison across systems, but not a precision-grade hallucination auditor.",
    "The intent-to-search-string mapping (Task 4.3.1) is a hand-built lookup table for the ~26 Bitext "
    "intents observed in this sample; a production system would need this maintained as the intent "
    "taxonomy evolves.",
])

h2("Recommendations & Future Work")
bullets([
    "Increase LoRA training steps/epochs and dataset size (toward the 5,000-row ceiling) once GPU "
    "compute is available, and re-run with QLoRA (4-bit) for a closer match to the reference T4 environment.",
    "Expand k in retrieval (top-2/3) for categories with higher SOP-to-SOP TF-IDF overlap (see "
    "Data_Understanding_and_EDA.ipynb, Task 1.3.2/1.3.3) to reduce single-document retrieval misses.",
    "Replace the heuristic hallucination detector with an LLM-judge or human-in-the-loop audit before "
    "any production deployment decision.",
    "Extend the intent-to-search-string mapping to a learned retrieval-routing layer instead of a "
    "static lookup table, so new intents do not require manual mapping updates.",
])

doc.save(OUT)
print("Wrote", OUT)
