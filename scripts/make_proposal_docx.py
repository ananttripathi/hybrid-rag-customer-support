"""Generates an editable Project_Proposal_and_Methodology.docx (same content as the PDF)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "Project_Proposal_and_Methodology.docx"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x0b, 0x3d, 0x91)

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x16, 0x32, 0x4f)

def p(text_runs):
    """text_runs: list of (text, bold, mono) tuples, or a plain string."""
    para = doc.add_paragraph()
    if isinstance(text_runs, str):
        text_runs = [(text_runs, False, False)]
    for text, bold, mono in text_runs:
        r = para.add_run(text)
        r.bold = bold
        if mono:
            r.font.name = "Courier New"
    return para

def bullets(items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        if isinstance(item, str):
            item = [(item, False, False)]
        for text, bold, mono in item:
            r = para.add_run(text)
            r.bold = bold
            if mono:
                r.font.name = "Courier New"

def mono_block(text):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()

# ---------------- Title ----------------
title("Project Proposal & Methodology")
subtitle("Hybrid RAG & Fine-Tuning for Customer Support")

# ---------------- 1. Project Scope ----------------
h1("1. Project Scope")
h2("1.1.1 Target Use Case")
p([("Domain: ", True, False),
   ("Automated Tier-1 customer support for the post-purchase / order-operations category of an "
    "e-commerce business. This scope was chosen because it is the single largest, most repetitive "
    "slice of inbound support volume, and it maps directly onto the corporate SOP corpus provided "
    "(refund, shipping, returns, account/password recovery, billing, subscriptions, escalation, "
    "order tracking, payment methods, data privacy, and working hours).", False, False)])
p([("In scope intents ", True, False),
   ("(drawn from the Bitext categories present in the sampled data): order status/tracking, "
    "cancellations, refunds, returns, shipping delays, password/account recovery, billing disputes, "
    "payment-method changes, subscription cancellation, and complaint escalation.", False, False)])
p([("Out of scope: ", True, False),
   ("pre-sales / product-discovery conversations, marketing content, and any action that mutates "
    "account or payment state (the system only classifies intent and drafts a policy-grounded "
    "response — it never executes a refund, cancellation, or account change itself).", False, False)])

h2("Expected JSON Interface")
p("Input to the fine-tuned intent router (Stage 4):")
mono_block('{"query": "the raw, possibly noisy customer message"}')
p("Output of the intent router — the structured target the model is fine-tuned to emit:")
mono_block('{"intent": "track_order", "category": "SHIPPING"}')
p("This structured output is what drives the Hybrid RAG retrieval step (Task 4.3): the intent is "
  "mapped to a specific SOP search string instead of the raw, noisy query.")

h2("1.1.2 Success Criteria")
p("Quantitative targets, defined before development began, evaluated on the held-out test split "
  "described in Stage 2:")
table(
    ["Metric", "Baseline (expected)", "Target for Hybrid RAG (V2)"],
    [
        ["Format Adherence Rate", "n/a (free text)", "≥ 90% valid JSON"],
        ["Intent Accuracy (exact match)", "n/a", "≥ 75% on test split, ≥ 60% on adversarial subset"],
        ["ROUGE-1 / ROUGE-L", "low (ungrounded)", "Improve over Baseline and V1 (Naive RAG)"],
        ["BLEU", "low (ungrounded)", "Improve over Baseline and V1 (Naive RAG)"],
        ["Consistency Rate", "100% (greedy)", "100% (greedy, do_sample=False)"],
        ["Hallucination Frequency", "high (unconstrained)", "≥ 40% relative reduction vs Baseline"],
    ],
)
p("These targets are re-measured identically for Baseline, Solution V1, and Solution V2 in the "
  "Comparative Analysis Report, so improvement is attributable and quantified stage-by-stage.")

doc.add_page_break()

# ---------------- 2. Proposed Methodology ----------------
h1("2. Proposed Methodology")

h2("1.4.1 Baseline Model Selection")
p([("Selected model: ", True, False), ("Qwen/Qwen2.5-1.5B-Instruct", False, True)])
bullets([
    [("Hardware constraints: ", True, False),
     ("the assignment targets a free-tier T4 GPU (Colab); this project was additionally executed "
      "end-to-end locally on Apple Silicon (M-series, MPS backend, no CUDA). At 1.5B parameters the "
      "model loads comfortably without requiring the 8B-class model's >16GB VRAM footprint.", False, False)],
    [("Dataset characteristics: ", True, False),
     ("the Bitext-derived instruction set (~1,500 rows after cleaning) is small; a 1.5B model has "
      "enough capacity to learn the JSON-extraction task via LoRA without severe overfitting risk.", False, False)],
    [("Fine-tuning feasibility: ", True, False),
     ("Qwen2.5-Instruct ships with a native ChatML template (apply_chat_template) and strong "
      "out-of-the-box structured-output behaviour, improving LoRA sample-efficiency for the JSON-router task.", False, False)],
    [("Rejected alternatives: ", True, False),
     ("TinyLlama-1.1B-Chat is lighter but empirically weaker at emitting strict JSON; "
      "Llama-3-8B-Instruct gives better raw quality but its memory footprint (>16GB even 4-bit) is "
      "infeasible on the constrained hardware used here.", False, False)],
    [("Hardware adaptation note: ", True, False),
     ("because execution is on Apple Silicon rather than a CUDA GPU, bitsandbytes 4-bit quantisation "
      "(CUDA-only) is unavailable. The model is instead loaded in fp16 directly on the MPS device — "
      "functionally equivalent for demonstrating the RAG + PEFT pipeline, and explicitly documented "
      "as a deviation from the Colab/T4 reference environment.", False, False)],
])

h2("1.4.2 Retrieval Strategy")
bullets([
    [("Embedding model: ", True, False),
     ("sentence-transformers/all-MiniLM-L6-v2 (384-dim, CPU-friendly, ~80MB) — used identically "
      "across NB4, NB5, and NB7 to keep retrieval results comparable across evaluation stages.", False, False)],
    [("Chunking: ", True, False),
     ("each corporate SOP Markdown file is treated as a single retrieval unit (documents are short, "
      "150-250 words each, single-topic) — no further sub-chunking required at this corpus size.", False, False)],
    [("Vector storage: ", True, False),
     ("ChromaDB, persisted to ./chroma_db, one collection for the full SOP corpus.", False, False)],
    [("Similarity search: ", True, False),
     ("naive RAG (V1) embeds the raw customer query and retrieves the top-1 nearest document (k=1). "
      "Hybrid RAG (V2) instead searches using the structured intent string extracted by the "
      "fine-tuned router, which is far less noisy than the raw query.", False, False)],
    [("Retrieval evaluation: ", True, False),
     ("validated directly (does the retrieved document match the query's true policy area?) and "
      "indirectly, via the downstream ROUGE/BLEU and hallucination-frequency scores of the generated answer.", False, False)],
])

h2("1.4.3 Fine-Tuning Strategy")
bullets([
    [("Technique: ", True, False),
     ("LoRA (not QLoRA) via peft — QLoRA's 4-bit quantisation path depends on bitsandbytes, which "
      "requires CUDA and is unavailable on the Apple Silicon execution environment used here.", False, False)],
    [("Target modules: ", True, False), ("q_proj, v_proj (attention projections) — updates ~0.5-1% of total parameters.", False, False)],
    [("Adapter config: ", True, False), ("rank r=16, lora_alpha=32 (2×r), lora_dropout=0.05, task_type=CAUSAL_LM.", False, False)],
    [("ChatML structure: ", True, False),
     ("system + user (raw query) + assistant, where the assistant turn is forced to a strict JSON "
      "string {\"intent\": ..., \"category\": ...} generated via json.dumps, with -100 label-masking "
      "applied to padding tokens so loss is computed only on real content.", False, False)],
    [("Training loop: ", True, False),
     ("custom PyTorch loop with AdamW (lr=2e-4), batch size 2, bf16 precision + gradient checkpointing "
      "(reduced from an initial fp32/batch-4 configuration after that caused system-wide memory swap "
      "thrashing on the 24GB execution machine), validation every VAL_EVERY steps, early stopping on "
      "validation loss (PATIENCE=3).", False, False)],
    [("Checkpointing / logging: ", True, False),
     ("best-val-loss checkpoint saved to ./intent_lora_best/, final checkpoint to ./intent_lora/, "
      "plus a training_log.csv and loss-curve plot as reproducibility evidence.", False, False)],
])

h2("1.4.4 Evaluation Framework")
bullets([
    [("Format Adherence Rate: ", True, False), ("percentage of router outputs that parse successfully with json.loads().", False, False)],
    [("Intent Accuracy: ", True, False),
     ("exact-match rate of the extracted intent field against ground truth, plus a fuzzy-match "
      "variant (e.g. order_tracking vs track_order) to capture near-misses.", False, False)],
    [("ROUGE-1 / ROUGE-L / BLEU: ", True, False),
     ("computed against SOP-grounded reference answers (not the generic Bitext response column), so "
      "policy-specific language is rewarded.", False, False)],
    [("Consistency Rate: ", True, False),
     ("repeated greedy-decoded generations (do_sample=False, temperature=None) for the same input "
      "must be identical.", False, False)],
    [("Hallucination Frequency: ", True, False),
     ("percentage of generations containing claims (dates, numbers, policy terms) not supported by "
      "the retrieved SOP context, checked for Baseline, V1, and V2 identically.", False, False)],
    [("All five metrics are re-computed under an identical configuration (same MODEL_ID, same "
      "embedding model, same test split) across all three system versions so the deltas are attributable.", False, False)],
])

doc.save(OUT)
print("Wrote", OUT)
